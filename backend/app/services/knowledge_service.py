import hashlib
import json
import re
import socket
from datetime import datetime
from ipaddress import ip_address
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

from fastapi import UploadFile

from app.core.config import settings
from app.db.session import rows_to_dicts
from app.services.auth_service import CurrentUser
from app.services.image_import_service import extract_text
from app.services.rag_service import normalize_community_name
from app.services.vector_service import delete_knowledge_chunks, upsert_knowledge_chunk


IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".webp"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
TEXT_EXTENSIONS = {".txt", ".md"}
PARENT_CHUNK_SIZE = 1800
CHILD_CHUNK_SIZE = 700
CHUNK_OVERLAP = 100
MAX_WEBPAGE_TEXT_CHARS = 300_000


def _safe_filename(filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    stem = re.sub(r"[^A-Za-z0-9_.-]+", "_", Path(filename).stem).strip("._") or "knowledge"
    digest = hashlib.sha1(f"{filename}-{datetime.now().isoformat()}".encode("utf-8")).hexdigest()[:10]
    return f"{stem}_{digest}{suffix}"


def _write_file(filename: str, content: bytes) -> Path:
    today = datetime.now().strftime("%Y%m%d")
    folder = settings.knowledge_root / today
    folder.mkdir(parents=True, exist_ok=True)
    path = folder / _safe_filename(filename)
    path.write_bytes(content)
    return path


def _extract_pdf_text(content: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:
        raise ValueError("当前环境缺少 pypdf，无法解析 PDF。") from exc

    reader = PdfReader(BytesIO(content))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"[第{index}页]\n{text}")
    return "\n\n".join(pages)


def _extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise ValueError("当前环境缺少 python-docx，无法解析 DOCX。") from exc

    document = Document(BytesIO(content))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_upload_text(file: UploadFile | None, file_content: bytes | None) -> tuple[str, str | None, str | None]:
    if file is None or not file.filename or not file_content:
        return "", None, None

    suffix = Path(file.filename).suffix.lower()
    file_path = _write_file(file.filename, file_content)

    if suffix in PDF_EXTENSIONS:
        text = _extract_pdf_text(file_content)
        if not text.strip():
            raise ValueError("PDF 未提取到文字。如为扫描件，请上传图片或可复制文字的 PDF。")
        return text, "pdf", str(file_path)
    if suffix in DOCX_EXTENSIONS:
        text = _extract_docx_text(file_content)
        if not text.strip():
            raise ValueError("DOCX 未提取到文字。")
        return text, "docx", str(file_path)
    if suffix in IMAGE_EXTENSIONS:
        text, error = extract_text(file_path)
        if error:
            raise ValueError(error)
        return text, "image", str(file_path)
    if suffix in TEXT_EXTENSIONS:
        return file_content.decode("utf-8", errors="ignore"), "text", str(file_path)
    raise ValueError("仅支持 PDF、DOCX、TXT、MD 和图片文件。")


def _clean_text(content: str) -> str:
    text = re.sub(r"\r\n?", "\n", content or "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _is_private_host(hostname: str) -> bool:
    if hostname in {"localhost", "0.0.0.0"} or hostname.endswith(".local"):
        return True
    try:
        addresses = {item[4][0] for item in socket.getaddrinfo(hostname, None)}
    except socket.gaierror as exc:
        raise ValueError("网页链接无法解析域名，请检查地址是否正确。") from exc
    for raw_address in addresses:
        try:
            address = ip_address(raw_address)
        except ValueError:
            continue
        if address.is_private or address.is_loopback or address.is_link_local or address.is_reserved:
            return True
    return False


def _extract_url_text(source_url: str | None) -> tuple[str, str | None, str | None]:
    url = (source_url or "").strip()
    if not url:
        return "", None, None

    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.hostname:
        raise ValueError("网页链接仅支持 http 或 https 地址。")
    if _is_private_host(parsed.hostname):
        raise ValueError("网页链接不能指向本机或内网地址。")

    try:
        from langchain_community.document_loaders import WebBaseLoader
    except Exception as exc:
        raise ValueError("当前环境缺少 langchain_community，无法加载网页链接。") from exc

    try:
        loader = WebBaseLoader(
            [url],
            header_template={
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/120.0.0.0 Safari/537.36"
                ),
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
            },
            requests_kwargs={"timeout": 12, "allow_redirects": False},
            raise_for_status=True,
            show_progress=False,
        )
        documents = loader.load()
    except ImportError as exc:
        raise ValueError("网页解析依赖缺失，请安装 beautifulsoup4 后重试。") from exc
    except Exception as exc:
        raise ValueError(f"网页链接读取失败：{exc}") from exc

    text = _clean_text("\n\n".join(doc.page_content for doc in documents if doc.page_content))
    if not text:
        raise ValueError("网页未提取到可用文字，请尝试粘贴正文或上传文件。")
    compact_text = re.sub(r"\s+", "", text)
    if len(compact_text) < 80 and ("\u672a\u77e5\u9519\u8bef" in compact_text or "\u8bf7\u7a0d\u540e\u518d\u8bd5" in compact_text):
        raise ValueError("网页返回了错误页，未读取到正文；请尝试复制正文上传，或换一个可公开访问的链接。")
    if compact_text in {"未知错误未知错误，请稍后再试", "未知错误请稍后再试"}:
        raise ValueError("网页返回了错误页，未读取到正文；请尝试复制正文上传，或换一个可公开访问的链接。")
    if len(text) > MAX_WEBPAGE_TEXT_CHARS:
        raise ValueError("网页正文过长，请改为上传文件或粘贴核心内容。")
    return text, "url", url


def _split_by_size(content: str, size: int, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = _clean_text(content)
    if not text:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + size)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(0, end - overlap)
    return chunks


def split_chunks(content: str) -> list[str]:
    return _split_by_size(content, CHILD_CHUNK_SIZE)


def _summary(content: str, limit: int = 180) -> str:
    text = re.sub(r"\s+", " ", content or "").strip()
    return text[:limit]


def _auto_tags(*values: str | None) -> str:
    tags: list[str] = []
    for value in values:
        for token in re.findall(r"[\u4e00-\u9fa5A-Za-z0-9]{2,}", value or ""):
            if token not in tags:
                tags.append(token)
    return ",".join(tags[:20])


def _metadata(
    *,
    document_id: int,
    title: str,
    city: str,
    community_name: str | None,
    knowledge_type: str,
    source_type: str,
    version: int,
    chunk_index: int,
    chunk_level: str,
    permission_scope: str,
    tags: str,
) -> dict[str, Any]:
    return {
        "document_id": document_id,
        "title": title,
        "city": city,
        "community_name": community_name or "",
        "knowledge_type": knowledge_type,
        "category": knowledge_type,
        "source_type": source_type,
        "version": version,
        "chunk_index": chunk_index,
        "chunk_level": chunk_level,
        "permission_scope": permission_scope,
        "tags": tags,
    }


def create_knowledge_document(
    conn,
    user: CurrentUser,
    *,
    title: str,
    community_name: str | None,
    knowledge_type: str,
    content: str | None,
    source_url: str | None,
    file: UploadFile | None,
    file_content: bytes | None,
) -> dict[str, Any]:
    title = title.strip()
    knowledge_type = knowledge_type.strip() or "楼盘信息"
    community_name = (community_name or "").strip() or None
    community_normalization = normalize_community_name(conn, user.city, community_name)
    if community_normalization.get("matched"):
        community_name = community_normalization["name"]
    typed_content = (content or "").strip()
    normalized_source_url = (source_url or "").strip() or None
    file_text, source_type, file_path = _extract_upload_text(file, file_content)
    url_text, url_source_type, normalized_source_url = _extract_url_text(normalized_source_url)
    source_type = source_type or url_source_type
    merged_content = _clean_text("\n\n".join(part for part in [typed_content, file_text, url_text] if part and part.strip()))
    if not title:
        raise ValueError("请填写知识标题。")
    if not merged_content:
        raise ValueError("请填写文字内容，或上传可识别的 PDF/DOCX/图片/文本文件。")

    scope_params = (user.city, community_name or "", knowledge_type)
    latest = conn.execute(
        "SELECT MAX(version) AS version FROM knowledge_documents WHERE city = ? AND COALESCE(community_name, '') = ? AND knowledge_type = ?",
        scope_params,
    ).fetchone()
    version = int((latest["version"] if latest else 0) or 0) + 1
    old_rows = conn.execute(
        "SELECT id FROM knowledge_documents WHERE city = ? AND COALESCE(community_name, '') = ? AND knowledge_type = ? AND status = 'active'",
        scope_params,
    ).fetchall()
    old_ids = [row["id"] for row in old_rows]
    if old_ids:
        placeholders = ",".join("?" for _ in old_ids)
        chroma_rows = conn.execute(
            f"SELECT chroma_id FROM knowledge_chunks WHERE chroma_id IS NOT NULL AND document_id IN ({placeholders})",
            tuple(old_ids),
        ).fetchall()
        delete_knowledge_chunks([row["chroma_id"] for row in chroma_rows if row["chroma_id"]])
        conn.execute(
            f"UPDATE knowledge_documents SET status = 'archived', archived_time = CURRENT_TIMESTAMP, modify_time = CURRENT_TIMESTAMP WHERE id IN ({placeholders})",
            tuple(old_ids),
        )
        conn.execute(f"UPDATE knowledge_chunks SET status = 'archived' WHERE document_id IN ({placeholders})", tuple(old_ids))

    tags = _auto_tags(title, community_name, knowledge_type, merged_content[:300])
    permission_scope = f"city:{user.city}"
    cursor = conn.execute(
        """
        INSERT INTO knowledge_documents(
            title, content, city, tags, status, community_name, knowledge_type, category,
            source_type, source_file, source_url, file_path, uploader_user_id, version, permission_scope,
            index_status, document_summary, indexed_time
        )
        VALUES (?, ?, ?, ?, 'active', ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'indexed', ?, CURRENT_TIMESTAMP)
        """,
        (
            title,
            merged_content,
            user.city,
            tags,
            community_name,
            knowledge_type,
            knowledge_type,
            source_type or ("url" if normalized_source_url else "text"),
            file.filename if file else normalized_source_url,
            normalized_source_url,
            file_path,
            user.id,
            version,
            permission_scope,
            _summary(merged_content, 400),
        ),
    )
    document_id = int(cursor.lastrowid)

    child_count = 0
    parent_chunks = _split_by_size(merged_content, PARENT_CHUNK_SIZE, overlap=0)
    for parent_index, parent in enumerate(parent_chunks):
        parent_meta = _metadata(
            document_id=document_id,
            title=title,
            city=user.city,
            community_name=community_name,
            knowledge_type=knowledge_type,
            source_type=source_type or "text",
            version=version,
            chunk_index=parent_index,
            chunk_level="parent",
            permission_scope=permission_scope,
            tags=tags,
        )
        parent_chroma_id = f"doc-{document_id}-parent-{parent_index}"
        parent_cursor = conn.execute(
            """
            INSERT INTO knowledge_chunks(
                document_id, city, community_name, knowledge_type, chunk_index, content,
                chunk_level, summary, tags, metadata_json, chroma_id
            )
            VALUES (?, ?, ?, ?, ?, ?, 'parent', ?, ?, ?, ?)
            """,
            (
                document_id,
                user.city,
                community_name,
                knowledge_type,
                parent_index,
                parent,
                _summary(parent),
                tags,
                json.dumps(parent_meta, ensure_ascii=False),
                parent_chroma_id,
            ),
        )
        parent_id = int(parent_cursor.lastrowid)
        upsert_knowledge_chunk(parent_chroma_id, _summary(parent, 600), parent_meta)

        for child_index, child in enumerate(_split_by_size(parent, CHILD_CHUNK_SIZE)):
            absolute_index = parent_index * 1000 + child_index
            child_meta = _metadata(
                document_id=document_id,
                title=title,
                city=user.city,
                community_name=community_name,
                knowledge_type=knowledge_type,
                source_type=source_type or "text",
                version=version,
                chunk_index=absolute_index,
                chunk_level="child",
                permission_scope=permission_scope,
                tags=tags,
            )
            chroma_id = f"doc-{document_id}-child-{absolute_index}"
            conn.execute(
                """
                INSERT INTO knowledge_chunks(
                    document_id, city, community_name, knowledge_type, chunk_index, content,
                    parent_chunk_id, chunk_level, summary, tags, metadata_json, chroma_id
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, 'child', ?, ?, ?, ?)
                """,
                (
                    document_id,
                    user.city,
                    community_name,
                    knowledge_type,
                    absolute_index,
                    child,
                    parent_id,
                    _summary(child),
                    tags,
                    json.dumps(child_meta, ensure_ascii=False),
                    chroma_id,
                ),
            )
            upsert_knowledge_chunk(chroma_id, child, child_meta)
            child_count += 1

    return {
        "id": document_id,
        "version": version,
        "chunks": child_count,
        "index_status": "indexed",
        "community_normalization": community_normalization,
    }


def list_knowledge_documents(conn, user: CurrentUser) -> list[dict]:
    where_sql = "1 = 1"
    params: tuple = ()
    if "admin" not in user.role_codes:
        where_sql = "kd.status = 'active' AND kd.city = ?"
        params = (user.city,)
    rows = conn.execute(
        f"""
        SELECT kd.id, kd.title, kd.city, kd.community_name, kd.knowledge_type, kd.category,
               kd.source_type, kd.source_file, kd.source_url, kd.version, kd.status, kd.index_status,
               kd.permission_scope, kd.create_time, kd.modify_time, kd.indexed_time,
               u.display_name AS uploader, kd.document_summary AS summary,
               COUNT(kc.id) AS chunk_count
        FROM knowledge_documents kd
        LEFT JOIN users u ON u.id = kd.uploader_user_id
        LEFT JOIN knowledge_chunks kc ON kc.document_id = kd.id AND kc.chunk_level = 'child'
        WHERE {where_sql}
        GROUP BY kd.id
        ORDER BY kd.id DESC
        LIMIT 200
        """,
        params,
    ).fetchall()
    return rows_to_dicts(rows)


def archive_knowledge_document(conn, document_id: int) -> bool:
    chroma_rows = conn.execute(
        "SELECT chroma_id FROM knowledge_chunks WHERE document_id = ? AND chroma_id IS NOT NULL",
        (document_id,),
    ).fetchall()
    delete_knowledge_chunks([row["chroma_id"] for row in chroma_rows if row["chroma_id"]])
    cursor = conn.execute(
        """
        UPDATE knowledge_documents
        SET status = 'archived', archived_time = CURRENT_TIMESTAMP, modify_time = CURRENT_TIMESTAMP
        WHERE id = ? AND status = 'active'
        """,
        (document_id,),
    )
    conn.execute("UPDATE knowledge_chunks SET status = 'archived' WHERE document_id = ?", (document_id,))
    return cursor.rowcount > 0
